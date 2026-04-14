import os
import shutil
from datetime import datetime
from docx import Document
from docxtpl import DocxTemplate
from docx.shared import Inches
from PIL import Image

# ─────────────────────────────────────────
# Config
# ─────────────────────────────────────────
TEMPLATES_DIR = "templates"
OUTPUTS_DIR   = "outputs"
TEMP_DIR      = "temp"

IMG_WIDTH_INCH = 3.0   # lebar foto dalam dokumen

# ─────────────────────────────────────────
# Helper: proses gambar (crop 4:3, resize)
# ─────────────────────────────────────────
def process_image(src_path, dst_path):
    img = Image.open(src_path).convert("RGB")
    w, h = img.size
    target = 4 / 3

    if w / h > target:
        new_w = int(h * target)
        left  = (w - new_w) // 2
        img   = img.crop((left, 0, left + new_w, h))
    else:
        new_h = int(w / target)
        top   = (h - new_h) // 2
        img   = img.crop((0, top, w, top + new_h))

    img = img.resize((960, 720))
    img.save(dst_path, "JPEG", quality=85)
    return dst_path


# ─────────────────────────────────────────
# Helper: get all paragraphs termasuk dalam tabel
# ─────────────────────────────────────────
def get_all_cells(doc):
    """Return semua cell dalam semua tabel di dokumen."""
    cells = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cells.append(cell)
    return cells


# ─────────────────────────────────────────
# Core: replace [IMG] dengan foto
# ─────────────────────────────────────────
def insert_images(doc, image_paths):
    """
    Replace setiap [IMG] placeholder dengan foto sesuai urutan.
    image_paths: list path foto, urutan harus sama dengan urutan [IMG] di template.
    """
    img_index = 0
    total_imgs = len(image_paths)

    for cell in get_all_cells(doc):
        for para in cell.paragraphs:
            if "[IMG]" not in para.text:
                continue

            if img_index >= total_imgs:
                raise ValueError(
                    f"Jumlah foto ({total_imgs}) kurang dari jumlah [IMG] di template"
                )

            img_path = image_paths[img_index]
            if not img_path or not os.path.exists(img_path):
                raise FileNotFoundError(
                    f"File foto tidak ditemukan: {img_path}"
                )

            # Hapus teks [IMG], insert gambar
            para.clear()
            run = para.add_run()
            run.add_picture(img_path, width=Inches(IMG_WIDTH_INCH))

            img_index += 1

    if img_index != total_imgs:
        raise ValueError(
            f"Jumlah foto ({total_imgs}) tidak sesuai dengan [IMG] di template ({img_index})"
        )

    return img_index


# ─────────────────────────────────────────
# Helper: strip [CAPTION] marker dari teks
# ─────────────────────────────────────────
def strip_captions(doc):
    """
    Hapus marker [CAPTION] dari semua cell di dokumen.
    Teks caption aslinya tetap ada, hanya prefix [CAPTION] yang dihapus.
    """
    for cell in get_all_cells(doc):
        for para in cell.paragraphs:
            if "[CAPTION]" not in para.text:
                continue

            # Ambil teks bersih tanpa [CAPTION]
            clean_text = para.text.replace("[CAPTION]", "").strip()

            # Simpan formatting run pertama (font, size, bold, dll)
            if para.runs:
                first_run = para.runs[0]
                font_name  = first_run.font.name
                font_size  = first_run.font.size
                font_bold  = first_run.font.bold
                font_color = first_run.font.color.rgb if first_run.font.color and first_run.font.color.type else None
            else:
                font_name = font_size = font_bold = font_color = None

            # Clear paragraph dan tulis ulang teks bersih
            para.clear()
            run = para.add_run(clean_text)

            # Restore formatting
            if font_name:  run.font.name  = font_name
            if font_size:  run.font.size  = font_size
            if font_bold:  run.font.bold  = font_bold
            if font_color:
                from docx.dml.color import ColorFormat
                run.font.color.rgb = font_color


# ─────────────────────────────────────────
# Core: generate dokumen
# ─────────────────────────────────────────
def generate_report(
    template_name,   # e.g. "monthlymvxr"
    tanggal,         # "11-04-2026"
    personil,        # "Misbahuddin"
    serial_number,   # "Hold Baggage Screening MVXR5000 Line 1 SN 6141187"
    image_paths,     # list path foto, urutan sesuai [IMG] di template
    output_filename=None
):
    """
    Generate laporan DOCX dari template.
    Returns: path ke file output DOCX.
    """
    os.makedirs(OUTPUTS_DIR, exist_ok=True)
    os.makedirs(TEMP_DIR,    exist_ok=True)

    # ── Step 1: Render teks dengan DocxTemplate ──────────────
    # Pastikan template_name tidak mengandung .docx (sudah di-strip di app.py)
    tpl_name  = template_name.replace(".docx", "")
    tpl_path  = os.path.join(TEMPLATES_DIR, f"{tpl_name}.docx")
    if not os.path.exists(tpl_path):
        raise FileNotFoundError(f"Template tidak ditemukan: {tpl_path}")

    temp_path = os.path.join(TEMP_DIR, f"temp_{tpl_name}_{os.getpid()}.docx")

    try:
        tpl = DocxTemplate(tpl_path)
        tpl.render({
            "date":          tanggal,
            "personil":      personil,
            "serial_number": serial_number
        })
        tpl.save(temp_path)
    except Exception as e:
        raise RuntimeError(f"Template render gagal: {e}")

    # ── Step 2: Proses & insert gambar ───────────────────────
    processed_paths = []
    try:
        for i, src in enumerate(image_paths):
            dst = os.path.join(TEMP_DIR, f"img_{os.getpid()}_{i}.jpg")
            process_image(src, dst)
            processed_paths.append(dst)

        doc = Document(temp_path)
        insert_images(doc, processed_paths)
        strip_captions(doc)  # Hapus marker [CAPTION] dari semua caption

        # ── Step 3: Simpan output ────────────────────────────
        if not output_filename:
            output_filename = (
                f"Dokumentasi_Laporan_PM_Peralatan_HBS_{tanggal}.docx"
            )
        output_path = os.path.join(OUTPUTS_DIR, output_filename)
        doc.save(output_path)

    finally:
        # Cleanup temp files
        for p in processed_paths:
            if os.path.exists(p):
                os.remove(p)
        if os.path.exists(temp_path):
            os.remove(temp_path)

    return output_path


# ─────────────────────────────────────────
# Helper: convert DOCX → PDF via LibreOffice
# ─────────────────────────────────────────
def convert_to_pdf(docx_path):
    """
    Convert DOCX ke PDF menggunakan LibreOffice.
    Returns: path ke file PDF.
    """
    import subprocess

    output_dir = os.path.dirname(docx_path)
    result = subprocess.run(
        [
            "libreoffice", "--headless",
            "--convert-to", "pdf",
            "--outdir", output_dir,
            docx_path
        ],
        capture_output=True, text=True, timeout=60
    )

    if result.returncode != 0:
        raise RuntimeError(f"LibreOffice convert gagal: {result.stderr}")

    pdf_path = docx_path.replace(".docx", ".pdf")
    if not os.path.exists(pdf_path):
        raise FileNotFoundError(f"PDF tidak ditemukan setelah convert: {pdf_path}")

    return pdf_path


# ─────────────────────────────────────────
# Test lokal
# ─────────────────────────────────────────
if __name__ == "__main__":
    import sys

    # Test dengan foto dummy
    test_images = [f"test_img_{i}.jpg" for i in range(29)]

    # Buat gambar test dummy
    os.makedirs("temp", exist_ok=True)
    for p in test_images:
        img = Image.new("RGB", (960, 720), color=(
            (i * 30) % 255,
            (i * 50) % 255,
            100
        ))
        img.save(p)

    try:
        out = generate_report(
            template_name  = "monthlymvxr",
            tanggal        = datetime.now().strftime("%d-%m-%Y"),
            personil       = "Test Personil",
            serial_number  = "MVXR5000 Line 1 SN 6141187",
            image_paths    = test_images,
            output_filename= "test_output.docx"
        )
        print(f"✅ Output: {out}")
    except Exception as e:
        print(f"❌ Error: {e}")
    finally:
        for p in test_images:
            if os.path.exists(p):
                os.remove(p)
