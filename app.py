import os
import uuid
import hmac
import hashlib
import logging
import threading
import time
from datetime import datetime
from collections import defaultdict
from dotenv import load_dotenv

# Load .env PERTAMA
load_dotenv()

from flask import Flask, request, jsonify, send_from_directory, render_template_string
import telebot
from generate import generate_report

# ─────────────────────────────────────────
# Config dari .env
# ─────────────────────────────────────────
BOT_TOKEN     = os.environ.get("BOT_TOKEN", "").strip()
CHAT_ID_GROUP = os.environ.get("CHAT_ID_GROUP", "").strip()
SERVER_URL    = os.environ.get("SERVER_URL", "").strip()

# Whitelist user ID Telegram yang boleh akses
# Format di .env: ALLOWED_USERS=123456789,987654321,555444333
_raw_users    = os.environ.get("ALLOWED_USERS", "").strip()
ALLOWED_USERS = set(u.strip() for u in _raw_users.split(",") if u.strip())

UPLOAD_DIR    = "uploads"
OUTPUTS_DIR   = "outputs"
TEMPLATES_DIR = "templates"

# Whitelist template yang diizinkan
ALLOWED_TEMPLATES = {
    "weeklymvxr", "weeklyrtt",
    "monthlymvxr", "monthlyrtt"
}

# Upload constraints
MAX_FILE_SIZE  = 5 * 1024 * 1024   # 5MB
ALLOWED_EXTS   = {"jpg", "jpeg", "png"}
ALLOWED_MIMES  = {"image/jpeg", "image/jpg", "image/png"}

# Rate limiting: max request per window
RATE_LIMIT_REQUESTS = 60   # max request
RATE_LIMIT_WINDOW   = 60   # per 60 detik

# Telegram initData max age (detik) — tolak jika > 1 jam
INIT_DATA_MAX_AGE = 3600

# ─────────────────────────────────────────
# Logging
# ─────────────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format="[%(asctime)s] %(levelname)s %(message)s",
    datefmt="%H:%M:%S"
)
log = logging.getLogger(__name__)

for d in [UPLOAD_DIR, OUTPUTS_DIR, "temp"]:
    os.makedirs(d, exist_ok=True)

# ─────────────────────────────────────────
# Rate Limiter
# ─────────────────────────────────────────
rate_store = defaultdict(list)  # ip → [timestamp, ...]

def is_rate_limited(identifier):
    now   = time.time()
    reqs  = rate_store[identifier]
    # Hapus request lama di luar window
    reqs  = [t for t in reqs if now - t < RATE_LIMIT_WINDOW]
    rate_store[identifier] = reqs
    if len(reqs) >= RATE_LIMIT_REQUESTS:
        return True
    rate_store[identifier].append(now)
    return False

def get_client_id():
    """Ambil identifier client: IP + user agent."""
    return request.remote_addr or "unknown"

# ─────────────────────────────────────────
# Telegram initData Validator
# ─────────────────────────────────────────
def verify_telegram_init_data(init_data_str):
    """
    Verifikasi initData dari Telegram WebApp menggunakan HMAC-SHA256.
    Returns: (valid: bool, user_id: str|None, error: str|None)
    """
    if not init_data_str:
        return False, None, "initData kosong"

    try:
        from urllib.parse import parse_qsl, unquote
        params = dict(parse_qsl(init_data_str, keep_blank_values=True))

        received_hash = params.pop("hash", None)
        if not received_hash:
            return False, None, "Hash tidak ada"

        # Cek expired (auth_date)
        auth_date = params.get("auth_date")
        if auth_date:
            age = int(time.time()) - int(auth_date)
            if age > INIT_DATA_MAX_AGE:
                return False, None, f"initData expired ({age}s)"

        # Build data_check_string
        data_check = "\n".join(
            f"{k}={v}" for k, v in sorted(params.items())
        )

        # HMAC verification
        secret_key = hmac.new(
            b"WebAppData",
            BOT_TOKEN.encode(),
            hashlib.sha256
        ).digest()

        expected_hash = hmac.new(
            secret_key,
            data_check.encode(),
            hashlib.sha256
        ).hexdigest()

        if not hmac.compare_digest(expected_hash, received_hash):
            return False, None, "Hash tidak valid"

        # Ambil user ID
        user_str = params.get("user", "{}")
        import json
        user_data = json.loads(unquote(user_str))
        user_id   = str(user_data.get("id", ""))

        return True, user_id, None

    except Exception as e:
        return False, None, f"Verifikasi error: {str(e)}"


def check_telegram_auth():
    """
    Decorator helper: validasi initData + whitelist user.
    Returns: (authorized: bool, user_id: str|None, error_response)
    """
    # Rate limit check
    client_id = get_client_id()
    if is_rate_limited(client_id):
        log.warning(f"Rate limit: {client_id}")
        return False, None, (jsonify({"success": False, "error": "Too many requests"}), 429)

    # Ambil initData dari header atau body
    init_data = (
        request.headers.get("X-Telegram-Init-Data") or
        request.form.get("initData") or
        request.args.get("initData") or
        (request.get_json(silent=True) or {}).get("initData", "")
    )

    # Jika tidak ada initData, tolak
    if not init_data:
        log.warning(f"No initData from {client_id}")
        return False, None, (jsonify({"success": False, "error": "Unauthorized"}), 403)

    valid, user_id, err = verify_telegram_init_data(init_data)
    if not valid:
        log.warning(f"Invalid initData from {client_id}: {err}")
        return False, None, (jsonify({"success": False, "error": "Unauthorized"}), 403)

    # Whitelist check
    if ALLOWED_USERS and user_id not in ALLOWED_USERS:
        log.warning(f"User {user_id} not in whitelist")
        return False, None, (jsonify({"success": False, "error": "Access denied"}), 403)

    return True, user_id, None


# ─────────────────────────────────────────
# Flask App
# ─────────────────────────────────────────
app = Flask(__name__, static_folder="miniapp")
app.config["MAX_CONTENT_LENGTH"] = MAX_FILE_SIZE

@app.after_request
def add_security_headers(response):
    response.headers["Access-Control-Allow-Origin"]  = "https://web.telegram.org"
    response.headers["Access-Control-Allow-Methods"] = "GET, POST, OPTIONS"
    response.headers["Access-Control-Allow-Headers"] = "Content-Type, X-Telegram-Init-Data"
    response.headers["X-Content-Type-Options"]       = "nosniff"
    response.headers["X-Frame-Options"]              = "SAMEORIGIN"
    response.headers["Referrer-Policy"]              = "strict-origin-when-cross-origin"
    return response

@app.route("/upload", methods=["OPTIONS"])
@app.route("/generate", methods=["OPTIONS"])
@app.route("/verify", methods=["OPTIONS"])
def options_handler():
    return "", 200

@app.route("/verify", methods=["POST"])
def verify_user_access():
    """Endpoint untuk mencek apakah user Telegram terdaftar di ALLOWED_USERS."""
    authorized, user_id, err_resp = check_telegram_auth()
    if not authorized:
        return err_resp
    return jsonify({"success": True})

bot = telebot.TeleBot(BOT_TOKEN)

# ─────────────────────────────────────────
# Halaman akses ditolak
# ─────────────────────────────────────────
DENIED_PAGE = """<!DOCTYPE html>
<html lang="id">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Akses Ditolak</title>
<style>
  * { box-sizing: border-box; margin: 0; padding: 0; }
  body {
    font-family: -apple-system, sans-serif;
    background: #1c1c1e; color: #f2f2f7;
    display: flex; align-items: center; justify-content: center;
    min-height: 100vh; padding: 20px;
  }
  .card {
    background: #2c2c2e; border-radius: 16px;
    padding: 32px 24px; text-align: center; max-width: 340px;
  }
  .icon { font-size: 48px; margin-bottom: 16px; }
  h1 { font-size: 20px; font-weight: 600; margin-bottom: 10px; }
  p  { font-size: 14px; color: rgba(242,242,247,0.6); line-height: 1.6; }
  .btn {
    display: inline-block; margin-top: 20px;
    padding: 12px 24px; background: #0a84ff;
    color: white; border-radius: 10px;
    font-size: 14px; font-weight: 600;
    text-decoration: none;
  }
</style>
</head>
<body>
<div class="card">
  <div class="icon">🔒</div>
  <h1>Akses Ditolak</h1>
  <p>Aplikasi ini hanya dapat diakses melalui Telegram.<br><br>
  Silakan buka melalui bot PM Report di Telegram.</p>
  <a class="btn" href="https://t.me/{{ bot_username }}">Buka di Telegram</a>
</div>
</body>
</html>"""

# ─────────────────────────────────────────
# Routes
# ─────────────────────────────────────────

@app.route("/")
def index():
    """Serve Mini App — hanya dari Telegram."""
    return send_from_directory("miniapp", "index.html")


@app.route("/upload", methods=["POST"])
def upload_photo():
    """Upload foto dengan validasi lengkap."""
    try:
        # Auth check
        authorized, user_id, err_resp = check_telegram_auth()
        if not authorized:
            return err_resp

        photo = request.files.get("photo")
        key   = request.form.get("key", "unknown")

        if not photo:
            return jsonify({"success": False, "error": "Tidak ada foto"}), 400

        # Validasi extension
        original_name = photo.filename or ""
        ext = original_name.rsplit(".", 1)[-1].lower() if "." in original_name else ""
        if ext not in ALLOWED_EXTS:
            return jsonify({"success": False, "error": "Format file tidak valid. Gunakan JPG atau PNG"}), 400

        # Validasi MIME type
        mime = photo.content_type or ""
        if mime not in ALLOWED_MIMES:
            return jsonify({"success": False, "error": "Tipe file tidak valid"}), 400

        # Validasi ukuran (double check)
        photo.seek(0, 2)
        size = photo.tell()
        photo.seek(0)
        if size > MAX_FILE_SIZE:
            return jsonify({"success": False, "error": f"Ukuran file terlalu besar (max 5MB)"}), 400

        # Simpan dengan nama UUID + timestamp (hindari overwrite)
        ts       = int(time.time())
        filename = f"{uuid.uuid4().hex}_{ts}.jpg"
        path     = os.path.join(UPLOAD_DIR, filename)
        photo.save(path)

        log.info(f"Upload OK: user={user_id} key={key} size={size//1024}KB → {filename}")
        return jsonify({"success": True, "filename": filename})

    except Exception as e:
        log.error(f"Upload error: {type(e).__name__}")
        return jsonify({"success": False, "error": "Upload gagal, coba lagi"}), 500


@app.route("/generate", methods=["POST"])
def generate():
    """Generate laporan dengan validasi penuh."""
    try:
        # Auth check
        authorized, user_id, err_resp = check_telegram_auth()
        if not authorized:
            return err_resp

        data = request.get_json()
        if not data:
            return jsonify({"success": False, "error": "Invalid request"}), 400

        template      = data.get("template", "").strip()
        tanggal       = data.get("tanggal", "").strip()
        personil      = data.get("personil", "").strip()
        serial_number = data.get("serial_number", "").strip()
        photos        = data.get("photos", {})
        chat_id       = data.get("chat_id")

        # Validasi field wajib
        if not all([template, tanggal, personil, serial_number]):
            return jsonify({"success": False, "error": "Data tidak lengkap"}), 400

        # Strip .docx dan validasi whitelist template
        template_clean = template.replace(".docx", "").strip().lower()
        if template_clean not in ALLOWED_TEMPLATES:
            log.warning(f"Template tidak valid: {template_clean} dari user={user_id}")
            return jsonify({"success": False, "error": "Template tidak valid"}), 400

        # Validasi karakter berbahaya di input
        for field_name, field_val in [("tanggal", tanggal), ("personil", personil)]:
            if any(c in field_val for c in ["../", "..\\", "<", ">", ";", "|"]):
                return jsonify({"success": False, "error": f"Input {field_name} tidak valid"}), 400

        # Susun urutan foto sesuai urutan [IMG] di template
        image_filenames = (
            photos.get("mvxr", []) +
            photos.get("rtt",  []) +
            photos.get("s3",   [])
        )

        # Validasi setiap filename foto
        image_paths = []
        for fname in image_filenames:
            if not fname:
                return jsonify({"success": False, "error": "Ada foto yang belum diupload"}), 400

            # Cegah path traversal
            safe_name = os.path.basename(fname)
            if safe_name != fname or ".." in fname or "/" in fname:
                log.warning(f"Path traversal attempt: {fname} dari user={user_id}")
                return jsonify({"success": False, "error": "Nama file tidak valid"}), 400

            # Validasi extension file upload
            if not any(fname.endswith(f".{e}") for e in ["jpg", "jpeg", "png"]):
                return jsonify({"success": False, "error": "Format file foto tidak valid"}), 400

            path = os.path.join(UPLOAD_DIR, safe_name)
            if not os.path.exists(path):
                return jsonify({"success": False, "error": "File foto tidak ditemukan"}), 400
            image_paths.append(path)

        log.info(f"Generate: user={user_id} template={template_clean} personil={personil} foto={len(image_paths)}")

        result = {}
        error  = {}

        def do_generate():
            try:
                tanggal_file    = tanggal.replace(" ", "_")
                output_filename = f"Dokumentasi_Laporan_PM_Peralatan_HBS_{tanggal_file}.docx"

                docx_path = generate_report(
                    template_name  = template_clean,
                    tanggal        = tanggal,
                    personil       = personil,
                    serial_number  = serial_number,
                    image_paths    = image_paths,
                    output_filename= output_filename
                )

                result["docx"] = docx_path
                send_to_telegram(docx_path, tanggal, personil, template_clean, chat_id)

                # Cleanup foto upload setelah generate
                for p in image_paths:
                    try:
                        if os.path.exists(p):
                            os.remove(p)
                    except:
                        pass

            except Exception as e:
                log.error(f"Generate error: {type(e).__name__}: {str(e)}")
                error["msg"] = str(e)

        t = threading.Thread(target=do_generate)
        t.start()
        t.join(timeout=120)

        if error:
            return jsonify({"success": False, "error": error["msg"]}), 500

        if not result:
            return jsonify({"success": False, "error": "Generate timeout, coba lagi"}), 500

        docx_filename = os.path.basename(result["docx"])
        return jsonify({
            "success":  True,
            "doc_url":  f"{SERVER_URL}/download/{docx_filename}",
        })

    except Exception as e:
        log.error(f"Generate route error: {type(e).__name__}")
        return jsonify({"success": False, "error": "Terjadi kesalahan, coba lagi"}), 500


@app.route("/download/<filename>")
def download_file(filename):
    """Download file output — hanya nama file yang aman."""
    try:
        # Auth check
        authorized, user_id, err_resp = check_telegram_auth()
        if not authorized:
            return err_resp

        safe_name = os.path.basename(filename)
        if safe_name != filename or ".." in filename:
            return jsonify({"success": False, "error": "Nama file tidak valid"}), 400

        if not filename.endswith(".docx"):
            return jsonify({"success": False, "error": "Tipe file tidak valid"}), 400

        return send_from_directory(OUTPUTS_DIR, safe_name, as_attachment=True)

    except Exception as e:
        log.error(f"Download error: {type(e).__name__}")
        return jsonify({"success": False, "error": "File tidak ditemukan"}), 404


@app.route("/captions/<template_name>")
def get_captions(template_name):
    """Extract caption dari template DOCX."""
    try:
        # Auth check
        authorized, user_id, err_resp = check_telegram_auth()
        if not authorized:
            return err_resp

        # Validasi whitelist template
        tpl_name = template_name.replace(".docx", "").strip().lower()
        if tpl_name not in ALLOWED_TEMPLATES:
            return jsonify({"success": False, "error": "Template tidak valid"}), 400

        from docx import Document as DocxDocument
        tpl_path = os.path.join(TEMPLATES_DIR, f"{tpl_name}.docx")

        if not os.path.exists(tpl_path):
            return jsonify({"success": False, "error": "Template tidak ditemukan"}), 404

        doc      = DocxDocument(tpl_path)
        captions = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        txt = para.text.strip()
                        if "[CAPTION]" in txt:
                            captions.append(txt.replace("[CAPTION]", "").strip())

        return jsonify({"success": True, "template": tpl_name, "captions": captions, "total": len(captions)})

    except Exception as e:
        log.error(f"Captions error: {type(e).__name__}")
        return jsonify({"success": False, "error": "Gagal membaca template"}), 500


@app.route("/health")
def health():
    """Health check — tidak expose info internal."""
    return jsonify({"status": "ok"})


# ─────────────────────────────────────────
# Auto-cleanup file lama
# ─────────────────────────────────────────
def cleanup_old_files():
    """Hapus file upload & output yang lebih dari 1 hari."""
    while True:
        try:
            time.sleep(3600)  # Jalankan setiap 1 jam
            now     = time.time()
            max_age = 86400   # 1 hari dalam detik

            for folder in [UPLOAD_DIR, OUTPUTS_DIR]:
                for fname in os.listdir(folder):
                    fpath = os.path.join(folder, fname)
                    if os.path.isfile(fpath):
                        age = now - os.path.getmtime(fpath)
                        if age > max_age:
                            os.remove(fpath)
                            log.info(f"Cleanup: hapus {fpath}")
        except Exception as e:
            log.error(f"Cleanup error: {e}")


# ─────────────────────────────────────────
# Telegram: kirim DOCX ke group
# ─────────────────────────────────────────
def send_to_telegram(docx_path, tanggal, personil, template, user_chat_id=None):
    """
    Kirim DOCX ke Telegram via HTTP API langsung (bukan via bot object).
    Lebih reliable di Render.com karena tidak bergantung pada bot polling instance.
    """
    import urllib.request
    import urllib.parse

    caption = (
        f"📋 *Laporan PM HBS*\n"
        f"📅 {tanggal}\n"
        f"👤 {personil}\n"
        f"📁 {template}\n"
        f"✅ Generated via PM Bot"
    )

    targets = [CHAT_ID_GROUP]
    if user_chat_id and str(user_chat_id) != str(CHAT_ID_GROUP):
        targets.append(str(user_chat_id))

    api_url = f"https://api.telegram.org/bot{BOT_TOKEN}/sendDocument"
    filename = os.path.basename(docx_path)

    for chat in targets:
        try:
            # Build multipart form-data secara manual
            boundary = f"----FormBoundary{uuid.uuid4().hex}"
            with open(docx_path, "rb") as f:
                file_data = f.read()

            body = (
                f"--{boundary}\r\n"
                f'Content-Disposition: form-data; name="chat_id"\r\n\r\n'
                f"{chat}\r\n"
                f"--{boundary}\r\n"
                f'Content-Disposition: form-data; name="caption"\r\n\r\n'
                f"{caption}\r\n"
                f"--{boundary}\r\n"
                f'Content-Disposition: form-data; name="parse_mode"\r\n\r\n'
                f"Markdown\r\n"
                f"--{boundary}\r\n"
                f'Content-Disposition: form-data; name="document"; filename="{filename}"\r\n'
                f"Content-Type: application/vnd.openxmlformats-officedocument.wordprocessingml.document\r\n\r\n"
            ).encode("utf-8") + file_data + f"\r\n--{boundary}--\r\n".encode("utf-8")

            req = urllib.request.Request(
                api_url,
                data=body,
                headers={"Content-Type": f"multipart/form-data; boundary={boundary}"}
            )
            with urllib.request.urlopen(req, timeout=60) as resp:
                result = resp.read().decode("utf-8")
                log.info(f"Telegram kirim OK → chat={chat} file={filename}")

        except Exception as e:
            log.error(f"Telegram send error ke {chat}: {type(e).__name__}: {str(e)}")

    log.info(f"Laporan selesai dikirim ke: {targets}")


# ─────────────────────────────────────────
# Telegram Bot Handlers
# ─────────────────────────────────────────
@bot.message_handler(commands=["start", "laporan"])
def cmd_laporan(message):
    user_id = str(message.from_user.id)
    if ALLOWED_USERS and user_id not in ALLOWED_USERS:
        bot.send_message(message.chat.id, "⛔ Akses ditolak.")
        log.warning(f"Bot: akses ditolak untuk user {user_id}")
        return

    markup = telebot.types.InlineKeyboardMarkup()
    markup.add(
        telebot.types.InlineKeyboardButton(
            text="📋 Buat Laporan PM",
            web_app=telebot.types.WebAppInfo(url=SERVER_URL)
        )
    )
    bot.send_message(
        message.chat.id,
        "🛡️ *PM Daily Report HBS*\n\nTap tombol di bawah untuk membuat laporan:",
        reply_markup=markup,
        parse_mode="Markdown"
    )


@bot.message_handler(commands=["status"])
def cmd_status(message):
    bot.send_message(message.chat.id, "✅ Bot aktif dan siap!")


@bot.message_handler(commands=["myid"])
def cmd_myid(message):
    """Helper untuk dapat user ID."""
    bot.send_message(message.chat.id, f"User ID kamu: `{message.from_user.id}`", parse_mode="Markdown")


@bot.message_handler(func=lambda m: True)
def fallback(message):
    bot.send_message(message.chat.id, "Ketik /laporan untuk membuat laporan PM.")


# ─────────────────────────────────────────
# Run
# ─────────────────────────────────────────
def run_bot():
    log.info("Bot polling started...")
    bot.infinity_polling(timeout=60, long_polling_timeout=60)


if __name__ == "__main__":
    # Validasi config wajib
    if not BOT_TOKEN:
        log.error("BOT_TOKEN tidak ada di .env!")
        exit(1)
    if not CHAT_ID_GROUP:
        log.error("CHAT_ID_GROUP tidak ada di .env!")
        exit(1)
    if not ALLOWED_USERS:
        log.warning("ALLOWED_USERS kosong — semua user Telegram bisa akses!")
    else:
        log.info(f"Whitelist aktif: {len(ALLOWED_USERS)} user diizinkan")

    import sys
    if len(sys.argv) > 1 and sys.argv[1] == "bot":
        # Mode Service Bot (dipanggil via systemd: python3 app.py bot)
        log.info("Running in BOT ONLY mode (for Gunicorn deployment)...")
        cleanup_thread = threading.Thread(target=cleanup_old_files, daemon=True)
        cleanup_thread.start()
        run_bot() # Menahan thread utama
    else:
        # Mode Local / Screen (menjalankan Flask built-in + Bot bersamaan)
        cleanup_thread = threading.Thread(target=cleanup_old_files, daemon=True)
        cleanup_thread.start()

        bot_thread = threading.Thread(target=run_bot, daemon=True)
        bot_thread.start()

        log.info(f"Flask running on port 5000 | URL: {SERVER_URL}")
        app.run(host="0.0.0.0", port=5000, debug=False)
